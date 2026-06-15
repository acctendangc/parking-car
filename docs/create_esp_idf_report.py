from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\esp_projects\parking-car\docs\Bao_cao_Parking_Car_ESP_IDF.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.add_run(value)
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run(" " + text)
    set_table_widths(table, [9360])
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "Báo cáo kỹ thuật | Parking Car ESP-IDF"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = GRAY

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Trang ")
    add_page_number(footer)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO ĐỀ TÀI")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("Hệ thống khóa giữ chỗ sạc thông minh cho xe điện\nSmart Parking Car for EV Charging")
    r.font.size = Pt(14)
    r.font.color.rgb = GRAY

    add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ["Nền tảng phần cứng", "ESP32 Dev Module"],
            ["Môi trường phát triển", "ESP-IDF dùng CMake, kết hợp Arduino core cho ESP32"],
            ["Code chính", r"C:\esp_projects\parking-car\firmware_idf\main\parking_lock_v2.cpp"],
            ["Chức năng chính", "Điều khiển khóa giữ chỗ, kiểm tra trạng thái sạc, cảnh báo chiếm chỗ/va chạm và giám sát qua Blynk"],
        ],
        [2300, 7060],
        LIGHT_BLUE,
    )

    add_heading(doc, "1. Giới thiệu đề tài", 1)
    add_para(doc, "Đề tài xây dựng mô hình hệ thống khóa giữ chỗ sạc thông minh cho xe điện. Hệ thống dùng ESP32 làm bộ điều khiển trung tâm, servo để mô phỏng cơ cấu khóa, công tắc để mô phỏng trạng thái cắm sạc, cảm biến khoảng cách VL53L1X để phát hiện vật thể/va chạm, LCD I2C để hiển thị trạng thái tại chỗ, LED và buzzer để cảnh báo.")
    add_para(doc, "Mục tiêu của việc chuyển từ Arduino IDE sang ESP-IDF là đưa project về cấu trúc rõ ràng hơn, dễ quản lý thư viện hơn và gần với cách tổ chức của các nhóm làm hệ thống nhúng chuyên nghiệp. Trong project hiện tại, code vẫn giữ phong cách Arduino thông qua Arduino core cho ESP32, nhưng quá trình build, cấu hình và nạp firmware được thực hiện bằng ESP-IDF.")

    add_heading(doc, "2. Mục tiêu và yêu cầu hệ thống", 1)
    add_bullets(doc, [
        "Cho phép điều khiển trạng thái hạ/nâng khóa giữ chỗ thông qua ứng dụng Blynk.",
        "Nhận biết xe đã cắm sạc bằng công tắc mô phỏng trạng thái charger.",
        "Cảnh báo nếu đã hạ khóa nhưng sau 10 giây chưa cắm sạc.",
        "Cảnh báo nếu đã rút sạc nhưng sau 10 giây vẫn chiếm chỗ.",
        "Cảnh báo va chạm hoặc vật thể quá gần bằng cảm biến khoảng cách VL53L1X.",
        "Hiển thị trạng thái tại chỗ bằng LCD 16x2, LED RGB, LED đơn và buzzer.",
        "Có thể build/nạp bằng ESP-IDF thay vì phụ thuộc hoàn toàn vào Arduino IDE.",
    ])

    add_heading(doc, "3. Tổng quan kiến trúc hệ thống", 1)
    add_para(doc, "ESP32 đóng vai trò bộ xử lý trung tâm. Các thiết bị ngoại vi giao tiếp với ESP32 qua nhiều dạng giao tiếp nhúng: GPIO số, PWM, I2C, Wi-Fi/Blynk và UART Serial debug.")
    add_table(
        doc,
        ["Khối", "Thành phần", "Vai trò"],
        [
            ["Điều khiển trung tâm", "ESP32 Dev Module", "Chạy chương trình chính, xử lý trạng thái, kết nối Wi-Fi/Blynk và điều khiển ngoại vi."],
            ["Cơ cấu khóa", "Servo SG90", "Mô phỏng thanh khóa giữ chỗ. Góc 0 độ là khóa nâng, góc 90 độ là hạ khóa cho xe vào."],
            ["Nhận biết sạc", "Công tắc KCD1", "Mô phỏng thao tác cắm/rút sạc. ESP32 đọc trạng thái qua GPIO input pull-up."],
            ["Cảnh báo", "Buzzer, LED RGB, LED đơn", "Buzzer báo lỗi, LED RGB báo trạng thái bãi/cảnh báo, LED đơn báo đang sạc."],
            ["Cảm biến", "VL53L1X", "Đo khoảng cách bằng giao tiếp I2C để phát hiện nguy cơ va chạm."],
            ["Hiển thị", "LCD I2C 16x2", "Hiển thị trạng thái hệ thống và thông báo cảnh báo tại chỗ."],
            ["Giám sát từ xa", "Blynk Cloud/App", "Nhận lệnh điều khiển khóa, gửi trạng thái khóa/sạc/khoảng cách lên ứng dụng."],
        ],
        [1800, 2200, 5360],
        LIGHT_BLUE,
    )

    add_heading(doc, "4. Cấu trúc thư mục project", 1)
    add_para(doc, "Thư mục làm việc chính của đề tài là C:\\esp_projects\\parking-car. Bên trong vẫn còn file .ino gốc để đối chiếu, nhưng bản dùng ESP-IDF nằm trong thư mục firmware_idf.")
    add_table(
        doc,
        ["Đường dẫn/thư mục", "Chức năng"],
        [
            [r"C:\esp_projects\parking-car", "Thư mục project chính, chứa mã nguồn, tài liệu, web demo và cấu hình Git."],
            [r"parking_lock_v2.ino", "Code Arduino gốc, dùng khi nạp trực tiếp bằng Arduino IDE."],
            [r"firmware_idf", "Project ESP-IDF dùng để build và flash bằng idf.py."],
            [r"firmware_idf\main", "Chứa code chính parking_lock_v2.cpp và CMakeLists.txt của component main."],
            [r"firmware_idf\components", "Chứa thư viện local như Blynk, ESP32Servo, LiquidCrystal_I2C, VL53L1X."],
            [r"firmware_idf\managed_components", "Chứa thư viện do ESP-IDF Component Manager tải/quản lý, quan trọng nhất là espressif__arduino-esp32."],
            [r"firmware_idf\build", "Thư mục sinh ra sau khi build, chứa file tạm và firmware .bin."],
            [r"firmware_idf\sdkconfig", "File cấu hình ESP-IDF của project."],
            [r"web", "Chứa file index.html, có thể dùng cho phần giao diện/demo web nếu cần mở rộng."],
            [r"docs", "Chứa tài liệu và báo cáo của đề tài."],
        ],
        [3100, 6260],
    )

    add_heading(doc, "5. Vì sao dùng ESP-IDF thay vì chỉ dùng Arduino IDE?", 1)
    add_para(doc, "Arduino IDE phù hợp khi cần viết và nạp nhanh một file .ino. Tuy nhiên, khi project có nhiều thư viện, nhiều module, cần cấu hình chi tiết hoặc cần chia sẻ cho nhóm, ESP-IDF giúp cấu trúc project rõ ràng và ổn định hơn.")
    add_table(
        doc,
        ["Tiêu chí", "Arduino IDE", "ESP-IDF trong project này"],
        [
            ["Cách build", "Bấm Verify/Upload trong IDE.", "Dùng idf.py build, idf.py flash, idf.py monitor."],
            ["Cấu trúc project", "Thường tập trung vào một file .ino.", "Tách main, components, managed_components, sdkconfig."],
            ["Quản lý thư viện", "Phụ thuộc thư viện cài trong Arduino Library Manager của máy.", "Thư viện được khai báo rõ trong CMakeLists và idf_component.yml."],
            ["Khả năng mở rộng", "Dễ bắt đầu nhưng khó quản lý khi project lớn.", "Phù hợp hơn cho project nhiều module, nhiều cấu hình."],
            ["Mức độ kỹ thuật", "Ẩn nhiều chi tiết build và cấu hình.", "Cho phép can thiệp sâu hơn vào cấu hình chip, FreeRTOS, flash, component."],
        ],
        [1700, 3630, 4030],
        LIGHT_BLUE,
    )
    add_callout(doc, "Lưu ý:", "Project hiện tại không bỏ hoàn toàn Arduino. Nó dùng Arduino core cho ESP32 bên trong ESP-IDF, vì vậy các hàm quen thuộc như setup(), loop(), pinMode(), digitalWrite(), millis() vẫn được dùng.")

    add_heading(doc, "6. Thiết kế phần cứng và phân công chân", 1)
    add_table(
        doc,
        ["Thiết bị", "GPIO", "Kiểu giao tiếp", "Chức năng"],
        [
            ["Servo SG90", "GPIO18", "PWM 50 Hz", "Điều khiển góc khóa: 0 độ nâng khóa, 90 độ hạ khóa."],
            ["Buzzer", "GPIO19", "Digital Output", "Bật/tắt theo chu kỳ 300 ms khi có cảnh báo."],
            ["Công tắc KCD1", "GPIO14", "Digital Input Pull-up", "Mô phỏng trạng thái cắm/rút sạc."],
            ["LED đơn", "GPIO15", "Digital Output", "Sáng khi hệ thống xác nhận đang sạc."],
            ["LED RGB đỏ", "GPIO26", "Digital Output", "Nhấp nháy đỏ khi cảnh báo."],
            ["LED RGB xanh lá", "GPIO27", "Digital Output", "Báo bãi trống/sẵn sàng."],
            ["LED RGB xanh dương", "GPIO25", "Digital Output", "Báo có xe đỗ/khóa đã hạ."],
            ["VL53L1X", "GPIO21 SDA, GPIO22 SCL", "I2C", "Đo khoảng cách, phát hiện vật thể quá gần."],
            ["LCD 16x2 I2C", "GPIO21 SDA, GPIO22 SCL", "I2C", "Hiển thị trạng thái hệ thống."],
        ],
        [1700, 1700, 2000, 3960],
        LIGHT_BLUE,
    )

    add_heading(doc, "7. Thiết kế giao tiếp nhúng", 1)
    add_para(doc, "Phần thiết kế giao tiếp nhúng mô tả cách ESP32 trao đổi dữ liệu với các ngoại vi và với ứng dụng giám sát. Trong hệ thống này, mỗi loại thiết bị dùng một kiểu giao tiếp phù hợp với vai trò của nó.")

    add_heading(doc, "7.1 Giao tiếp GPIO số", 2)
    add_para(doc, "GPIO được dùng cho các tín hiệu bật/tắt đơn giản. Công tắc KCD1 được cấu hình INPUT_PULLUP, nghĩa là khi công tắc đóng thì chân đọc mức LOW. Các LED và buzzer được cấu hình OUTPUT, ESP32 xuất mức HIGH/LOW để bật hoặc tắt thiết bị.")
    add_bullets(doc, [
        "Ưu điểm: đơn giản, phản hồi nhanh, phù hợp với tín hiệu nhị phân.",
        "Ứng dụng trong đề tài: đọc trạng thái công tắc sạc, bật LED đơn, điều khiển buzzer và từng kênh màu của LED RGB.",
        "Lưu ý thiết kế: công tắc dùng pull-up nên logic bị đảo; trong code, trạng thái bật được kiểm tra bằng digitalRead(switchPin) == LOW.",
    ])

    add_heading(doc, "7.2 Giao tiếp PWM với servo", 2)
    add_para(doc, "Servo SG90 không chỉ nhận HIGH/LOW mà cần xung PWM có chu kỳ phù hợp. Trong project, ESP32Servo cấu hình servo ở tần số 50 Hz, sau đó ghi góc bằng myServo.write(). Góc 0 độ biểu diễn trạng thái khóa nâng, góc 90 độ biểu diễn trạng thái hạ khóa cho xe vào.")
    add_bullets(doc, [
        "Tín hiệu điều khiển: PWM trên GPIO18.",
        "Trạng thái mặc định: angleDefault = 0 độ.",
        "Trạng thái mở/hạ khóa: angleUnlock = 90 độ.",
    ])

    add_heading(doc, "7.3 Giao tiếp I2C dùng chung", 2)
    add_para(doc, "ESP32 dùng chung bus I2C trên GPIO21 (SDA) và GPIO22 (SCL) cho hai thiết bị: cảm biến VL53L1X và LCD 16x2 I2C. I2C cho phép nhiều thiết bị cùng dùng hai dây tín hiệu, miễn là mỗi thiết bị có địa chỉ khác nhau.")
    add_table(
        doc,
        ["Thiết bị I2C", "Vai trò", "Thông tin trong code"],
        [
            ["VL53L1X", "Đọc khoảng cách để cảnh báo va chạm.", "Khởi tạo Short mode, timing budget 50000 us, đọc liên tục mỗi 50 ms."],
            ["LCD I2C 16x2", "Hiển thị trạng thái tại chỗ.", "Địa chỉ 0x3F, cập nhật nội dung khi trạng thái thay đổi để tránh nháy màn hình."],
        ],
        [1800, 3300, 4260],
    )
    add_para(doc, "Ngưỡng cảnh báo va chạm được đặt là 40 mm. Trong vòng lặp chính, ESP32 đọc khoảng cách khoảng mỗi 100 ms. Nếu khoảng cách nhỏ hơn ngưỡng, hệ thống bật cảnh báo va chạm, buzzer kêu và LED RGB đỏ nhấp nháy.")

    add_heading(doc, "7.4 Giao tiếp Wi-Fi và Blynk", 2)
    add_para(doc, "ESP32 kết nối Wi-Fi và Blynk Cloud để nhận lệnh điều khiển cũng như gửi trạng thái lên ứng dụng. Đây là phần giao tiếp ở tầng mạng, khác với GPIO/I2C/PWM là giao tiếp cục bộ với phần cứng.")
    add_table(
        doc,
        ["Virtual Pin", "Chiều dữ liệu", "Ý nghĩa"],
        [
            ["V9", "App -> ESP32", "Nút điều khiển hạ/nâng khóa. Khi V9 = 1, servo hạ khóa; khi V9 = 0, servo nâng khóa."],
            ["V6", "ESP32 -> App", "Trạng thái khóa: đã hạ hoặc đã nâng."],
            ["V7", "ESP32 -> App", "Trạng thái sạc: đang sạc hoặc không sạc."],
            ["V8", "ESP32 -> App", "Giá trị khoảng cách đo được từ VL53L1X."],
        ],
        [1600, 2200, 5560],
        LIGHT_BLUE,
    )

    add_heading(doc, "7.5 UART Serial debug", 2)
    add_para(doc, "Serial ở tốc độ 115200 baud được dùng để in log khởi động, trạng thái cảm biến và các cảnh báo. Đây là kênh kiểm tra quan trọng khi nạp firmware hoặc debug bằng idf.py monitor.")

    add_heading(doc, "7.6 Giao tiếp giữa Arduino core và ESP-IDF", 2)
    add_para(doc, "Do project dùng Arduino core trong ESP-IDF, chương trình có thêm hàm app_main(). Đây là điểm vào chuẩn của ESP-IDF. Bên trong app_main(), chương trình gọi initArduino(), sau đó gọi setup() một lần và gọi loop() lặp liên tục. Mỗi vòng lặp có vTaskDelay(pdMS_TO_TICKS(1)) để nhường thời gian cho FreeRTOS.")
    add_callout(doc, "Ý nghĩa:", "Cách này giúp giữ được phong cách lập trình Arduino nhưng vẫn build/nạp bằng ESP-IDF. Đây là lựa chọn phù hợp khi muốn chuyển project .ino sang ESP-IDF mà không viết lại toàn bộ theo API ESP-IDF thuần.")

    add_heading(doc, "8. Luồng hoạt động của hệ thống", 1)
    add_numbers(doc, [
        "Khi khởi động, ESP32 cấu hình GPIO, I2C, LCD, cảm biến VL53L1X, servo và kết nối Blynk.",
        "Trạng thái mặc định là bãi trống, servo ở góc 0 độ, LED RGB xanh lá.",
        "Người dùng bấm nút trên Blynk để hạ khóa. ESP32 điều khiển servo về 90 độ và bắt đầu đếm thời gian.",
        "Nếu trong 10 giây người dùng bật công tắc sạc, hệ thống xác nhận đang sạc, bật LED đơn và gửi trạng thái lên Blynk.",
        "Nếu sau 10 giây chưa cắm sạc, hệ thống bật cảnh báo: buzzer kêu, LED RGB đỏ nhấp nháy và LCD hiển thị cảnh báo.",
        "Nếu rút sạc khi xe vẫn đang chiếm chỗ, hệ thống bắt đầu đếm 10 giây. Hết thời gian mà chưa rời bãi thì cảnh báo chiếm chỗ.",
        "Trong toàn bộ quá trình, cảm biến VL53L1X liên tục kiểm tra khoảng cách. Nếu vật thể quá gần, cảnh báo va chạm được ưu tiên hiển thị.",
    ])

    add_heading(doc, "9. Cấu trúc build bằng ESP-IDF", 1)
    add_para(doc, "Project ESP-IDF được build từ thư mục firmware_idf. File CMakeLists.txt ở cấp project khai báo project cho ESP-IDF, còn main/CMakeLists.txt khai báo file nguồn và các component cần dùng.")
    add_table(
        doc,
        ["File", "Vai trò"],
        [
            [r"firmware_idf\CMakeLists.txt", "Khai báo project ESP-IDF và include hệ build project.cmake."],
            [r"firmware_idf\main\CMakeLists.txt", "Đăng ký file parking_lock_v2.cpp và yêu cầu các component: arduino-esp32, Blynk, ESP32Servo, VL53L1X, LiquidCrystal_I2C."],
            [r"firmware_idf\main\idf_component.yml", "Khai báo dependency espressif/arduino-esp32 để dùng Arduino core trong ESP-IDF."],
            [r"firmware_idf\sdkconfig.defaults", "Cấu hình mặc định, hiện có CONFIG_FREERTOS_HZ=1000."],
        ],
        [3300, 6060],
    )
    add_para(doc, "Các lệnh thường dùng:", bold_prefix=None)
    add_table(
        doc,
        ["Lệnh", "Chức năng"],
        [
            [r"cd C:\esp_projects\parking-car\firmware_idf", "Chuyển vào thư mục ESP-IDF project."],
            ["idf.py build", "Biên dịch project và sinh firmware trong thư mục build."],
            ["idf.py -p COM3 flash", "Nạp firmware vào ESP32 qua cổng COM tương ứng."],
            ["idf.py -p COM3 monitor", "Mở Serial Monitor để xem log từ ESP32."],
            ["idf.py -p COM3 flash monitor", "Nạp firmware rồi mở monitor ngay sau đó."],
        ],
        [3500, 5860],
        LIGHT_BLUE,
    )

    add_heading(doc, "10. Đánh giá ưu điểm và hạn chế", 1)
    add_table(
        doc,
        ["Nội dung", "Đánh giá"],
        [
            ["Ưu điểm", "Cấu trúc project rõ ràng, giữ được code Arduino quen thuộc, có giám sát Blynk, có cảnh báo đa kênh và dùng được nhiều kiểu giao tiếp nhúng."],
            ["Hạn chế", "Thông tin Wi-Fi/Blynk đang ghi trực tiếp trong code; mô hình dùng công tắc để mô phỏng cắm sạc nên chưa phải đo dòng/áp sạc thực tế; cơ cấu servo mới là mô phỏng khóa."],
            ["Hướng phát triển", "Tách thông tin Wi-Fi/token ra cấu hình riêng, bổ sung cảm biến dòng/áp, thêm trạng thái xác thực người dùng, lưu log và thiết kế vỏ cơ khí chắc chắn hơn."],
        ],
        [1800, 7560],
    )

    add_heading(doc, "11. Kết luận", 1)
    add_para(doc, "Hệ thống Parking Car cho thấy cách kết hợp giữa điều khiển nhúng cục bộ và giám sát từ xa. ESP32 vừa điều khiển phần cứng như servo, LED, buzzer, LCD, cảm biến khoảng cách, vừa giao tiếp với Blynk qua Wi-Fi để người dùng điều khiển và theo dõi trạng thái.")
    add_para(doc, "Việc chuyển từ Arduino IDE sang ESP-IDF không làm thay đổi mục tiêu vận hành của hệ thống, nhưng giúp project có cấu trúc chuyên nghiệp hơn. Với cách dùng Arduino core trong ESP-IDF, đề tài vẫn tận dụng được code Arduino sẵn có trong khi có thêm lợi ích của hệ build ESP-IDF.")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
